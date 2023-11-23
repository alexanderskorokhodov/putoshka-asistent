from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt


class DocsFileDataSource:

    def __init__(self, document):
        self.doc = document

    def makeShortContent(self, shortContentItems):
        table = self.doc.add_table(rows=len(shortContentItems), cols=2)

        # Set the width of each cell to half of the document width
        table.autofit = False
        table.columns[0].width = Inches(10)  # Adjust the width as needed
        table.columns[1].width = Inches(10)  # Adjust the width as needed

        for i in range(len(shortContentItems)):
            cell1 = table.cell(i, 0)
            cell1.text = shortContentItems[i][0]

            for paragraph in cell1.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(12)
                    run.font.name = "IBM Plex Sans"
                    run.font.bold = True

            cell2 = table.cell(i, 1)
            cell2.text = shortContentItems[i][1]
            for paragraph in cell2.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(12)
                    run.font.name = "IBM Plex Sans"
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    def makeSubtitle(self, text):
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run(text)
        font = run.font
        font.size = Pt(16)
        font.name = 'IBM Plex Sans'  # Set the font name
        font.bold = True

    def makeTermsSection(self, terms):
        self.makeSubtitle("Термины, используемые в лекции")

        for item in terms:
            print(item)
            paragraph = self.doc.add_paragraph()
            run = paragraph.add_run(item[0] + " ")
            font = run.font
            font.size = Pt(12)
            font.name = 'IBM Plex Sans'  # Set the font name
            font.bold = True

            run = paragraph.add_run("- " + item[1])
            font = run.font
            font.size = Pt(12)
            font.name = 'IBM Plex Sans'  # Set the font name
            font.bold = False



    def makeHeader(self, text):
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run(text)
        font = run.font
        font.size = Pt(22)
        font.name = 'IBM Plex Sans'  # Set the font name
        font.bold = True

    def makeTextWithBoldWords(self, words, paragraph):
        for item in words:
            run = paragraph.add_run(item[0])
            font = run.font
            font.size = Pt(12)  # Set the font size
            font.name = 'IBM Plex Sans'  # Set the font name
            font.bold = item[1]

    def makeLectureDocs(self, text, terms):

        self.makeHeader("Оглавление")
        shortContests = [[item[0], "0"] for item in text]
        self.makeShortContent(shortContests)

        self.makeTermsSection(terms=terms)

        for section in text:
            self.makeHeader(section[0])
            paragraph = self.doc.add_paragraph()
            self.makeTextWithBoldWords(section[1], paragraph)

    def getDocs(self):
        self.doc.save("demo.docx")