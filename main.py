from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt

doc = Document()

'''
header_font_style = doc.styles
font_charstyle = header_font_style.add_style('HeaderFontStyle', WD_STYLE_TYPE.CHARACTER)
header_font_object = font_charstyle.font
header_font_object.size = Pt(22)
header_font_object.name = 'IBM Plex Sans'



short_content_parag = doc.add_paragraph("")
short_content_parag.add_run("Оглавление", style='HeaderFontStyle').bold = True

'''



def makeShortContent(shortContentItems):
    table = doc.add_table(rows=len(shortContentItems), cols=2)

    # Set the width of each cell to half of the document width
    table.autofit = False
    table.columns[0].width = Inches(10)  # Adjust the width as needed
    table.columns[1].width = Inches(10)  # Adjust the width as needed

    # Access the first cell and add the first word
    for i in range(len(shortContentItems)):
        cell1 = table.cell(i, 0)
        cell1.text = shortContentItems[i][0]

        for paragraph in cell1.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(12)  # Adjust the font size as needed
                run.font.name = "IBM Plex Sans"
                run.font.bold = True  # Make the  bold

        cell2 = table.cell(i, 1)
        cell2.text = shortContentItems[i][1]
        for paragraph in cell2.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(12)  # Adjust the font size as needed
                run.font.name = "IBM Plex Sans"
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT


def makeHeader(text):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    font = run.font
    font.size = Pt(22)
    font.name = 'IBM Plex Sans'  # Set the font name
    font.bold = True


def createListSentences(sentences):
    for item in sentences:
        paragraph = doc.add_paragraph(style='List Bullet')
        makeTextWithBoldWords(item, paragraph)


def makeParagraph(text):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    font = run.font
    font.size = Pt(12)  # Set the font size
    font.name = 'IBM Plex Sans'  # Set the font name



def makeTextWithBoldWords(words, paragraph):
    for item in words:
        run = paragraph.add_run(item[0])
        font = run.font
        font.size = Pt(12)  # Set the font size
        font.name = 'IBM Plex Sans'  # Set the font name
        font.bold = item[1]



makeHeader("Оглавление")

makeShortContent(
    [
        ["Введение", "3"],
        ["Термины, используемые в лекции", "3"],
        ["Знакомство", "4"],
        ["План курса", "4"],
        ["1. Почему именно С#?", "5"],
        ["1. Почему именно С#?", "6"],
        ["1. Почему именно С#?", "7"],
        [" 4.1. Пример 1", "11"],
        ["1. Почему именно С#?", "12"],
        ["1. Почему именно С#?", "13"]
    ]
)

makeHeader("Введение")
makeParagraph("Этот курс")
createListSentences(
    [
        [["познакомит ", True], ["вас ", False], ["с ", False], ["синтаксисом ", False], ["языка ", False], ["программирования ", False], ["высокого ", False], ["уровня ", False]],
        [["познакомит ", True], ["вас ", False], ["с ", False], ["синтаксисом ", False], ["языка ", False], ["программирования ", False], ["высокого ", False], ["уровня ", False]],
        [["познакомит ", True], ["вас ", False], ["с ", False], ["синтаксисом ", False], ["языка ", False], ["программирования ", False], ["высокого ", False], ["уровня ", False]]
    ]
)

paragraph = doc.add_paragraph()
makeTextWithBoldWords(
    [
        ["даст ", True],
        ["вам ", False],
        ["навыки ", True],
        ["решения ", False],
        ["базовых ", False],
        ["алгоритмических ", False],
        ["задач ", False]
    ],
    paragraph
)


