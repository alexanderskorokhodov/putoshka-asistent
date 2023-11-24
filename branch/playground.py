from docx.oxml import parse_xml, register_element_cls


from docx import Document
from docx.shared import Inches, Pt, RGBColor


# refer to docx.oxml.shape.CT_Inline

document = Document()

def createTitlePage(header, subtitle, img_src):
    register_element_cls('wp:anchor', CT_Anchor)
    p = document.add_paragraph()
    add_float_picture(p, img_src, width=Inches(15.0), pos_x=Pt(0), pos_y=Pt(0))
    paragraph = document.add_heading(header)
    run = paragraph.runs[0]
    font = run.font
    font.color.rgb = RGBColor(255, 255, 255)  # font.size = Pt(35)
    font.name = 'IBM Plex Sans'  # Set the font name
    font.bold = True
    font.size = Pt(45)
    paragraph.paragraph_format.space_before = Pt(100)  # Размер отступа в точках

    paragraph = document.add_heading()
    run = paragraph.add_run(subtitle)
    font = run.font
    font.size = Pt(30)
    font.name = 'IBM Plex Sans'  # Set the font name
    font.bold = True
    font.color.rgb = RGBColor(100, 75, 135)    # font.size = Pt(35)
    paragraph.paragraph_format.space_before = Pt(20)  # Размер отступа в точках

def makeHeader(text):
    paragraph = document.add_heading()
    run = paragraph.add_run(text)
    font = run.font
    font.size = Pt(22)
    font.name = 'IBM Plex Sans'  # Set the font name
    font.bold = True

createTitlePage("Введение", "Знакомство с языками программирования", "../doc_back.png")
document.add_page_break()
makeHeader("Введение")

document.save('output.docx')