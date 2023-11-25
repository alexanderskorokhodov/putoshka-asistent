
from docx.shared import Inches, Pt, RGBColor
from docx2pdf import convert
from docx.oxml.ns import nsdecls
from docx.oxml.shape import CT_Picture
from docx.oxml.xmlchemy import BaseOxmlElement, OneAndOnlyOne
from docx.oxml import parse_xml, register_element_cls
class CT_Anchor(BaseOxmlElement):
    """
    ``<w:anchor>`` element, container for a floating image.
    """
    extent = OneAndOnlyOne('wp:extent')
    docPr = OneAndOnlyOne('wp:docPr')
    graphic = OneAndOnlyOne('a:graphic')

    @classmethod
    def new(cls, cx, cy, shape_id, pic, pos_x, pos_y):
        """
        Return a new ``<wp:anchor>`` element populated with the values passed
        as parameters.
        """
        anchor = parse_xml(cls._anchor_xml(pos_x, pos_y))
        anchor.extent.cx = cx
        anchor.extent.cy = cy
        anchor.docPr.id = shape_id
        anchor.docPr.name = 'Picture %d' % shape_id
        anchor.graphic.graphicData.uri = (
            'http://schemas.openxmlformats.org/drawingml/2006/picture'
        )
        anchor.graphic.graphicData._insert_pic(pic)
        return anchor

    @classmethod
    def new_pic_anchor(cls, shape_id, rId, filename, cx, cy, pos_x, pos_y):
        """
        Return a new `wp:anchor` element containing the `pic:pic` element
        specified by the argument values.
        """
        pic_id = 0  # Word doesn't seem to use this, but does not omit it
        pic = CT_Picture.new(pic_id, filename, rId, cx, cy)
        anchor = cls.new(cx, cy, shape_id, pic, pos_x, pos_y)
        anchor.graphic.graphicData._insert_pic(pic)
        return anchor

    @classmethod
    def _anchor_xml(cls, pos_x, pos_y):
        return (
            '<wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="0" \n'
            '           behindDoc="1" locked="0" layoutInCell="1" allowOverlap="1" \n'
            '           %s>\n'
            '  <wp:simplePos x="0" y="0"/>\n'
            '  <wp:positionH relativeFrom="page">\n'
            '    <wp:posOffset>%d</wp:posOffset>\n'
            '  </wp:positionH>\n'
            '  <wp:positionV relativeFrom="page">\n'
            '    <wp:posOffset>%d</wp:posOffset>\n'
            '  </wp:positionV>\n'                    
            '  <wp:extent cx="914400" cy="914400"/>\n'
            '  <wp:wrapNone/>\n'
            '  <wp:docPr id="666" name="unnamed"/>\n'
            '  <wp:cNvGraphicFramePr>\n'
            '    <a:graphicFrameLocks noChangeAspect="1"/>\n'
            '  </wp:cNvGraphicFramePr>\n'
            '  <a:graphic>\n'
            '    <a:graphicData uri="URI not set"/>\n'
            '  </a:graphic>\n'
            '</wp:anchor>' % ( nsdecls('wp', 'a', 'pic', 'r'), int(pos_x), int(pos_y) )
        )

class LectureDocxDataSource:

    def __init__(self, document):
        self.doc = document

    def new_pic_anchor(self, part, image_descriptor, width, height, pos_x, pos_y):
        """Return a newly-created `w:anchor` element.

        The element contains the image specified by *image_descriptor* and is scaled
        based on the values of *width* and *height*.
        """
        rId, image = part.get_or_add_image(image_descriptor)
        cx, cy = image.scaled_dimensions(width, height)
        shape_id, filename = part.next_id, image.filename
        return CT_Anchor.new_pic_anchor(shape_id, rId, filename, cx, cy, pos_x, pos_y)

    def createTitlePage(self, header, subtitle, img_src):
        register_element_cls('wp:anchor', CT_Anchor)
        p = self.doc.add_paragraph()
        self.add_float_picture(p, img_src, width=Inches(15.0), pos_x=Pt(0), pos_y=Pt(0))
        paragraph = self.doc.add_heading(header)
        run = paragraph.runs[0]
        font = run.font
        font.color.rgb = RGBColor(255, 255, 255)  # font.size = Pt(35)
        font.name = 'IBM Plex Sans'  # Set the font name
        font.bold = True
        font.size = Pt(45)
        paragraph.paragraph_format.space_before = Pt(100)  # Размер отступа в точках

        paragraph = self.doc.add_heading()
        run = paragraph.add_run(subtitle)
        font = run.font
        font.size = Pt(30)
        font.name = 'IBM Plex Sans'  # Set the font name
        font.bold = True
        font.color.rgb = RGBColor(100, 75, 135)  # font.size = Pt(35)
        paragraph.paragraph_format.space_before = Pt(20)  # Размер отступа в точках

    # refer to docx.text.run.add_picture
    def add_float_picture(self, p, image_path_or_stream, width=None, height=None, pos_x=0, pos_y=0):
        """Add float picture at fixed position `pos_x` and `pos_y` to the top-left point of page.
        """
        run = p.add_run()
        anchor = self.new_pic_anchor(run.part, image_path_or_stream, width, height, pos_x, pos_y)
        run._r.add_drawing(anchor)

    def makeShortContent(self, shortContentItems):

        # Set the width of each cell to half of the document width
        for i in range(len(shortContentItems)):
            paragraph = self.doc.add_paragraph()
            run = paragraph.add_run(f"{i+1}. {shortContentItems[i][0]}")
            font = run.font
            font.size = Pt(14)  # Set the font size
            font.name = 'IBM Plex Sans'  # Set the font name
            font.bold = True
            paragraph.paragraph_format.space_before = Pt(15)  # Размер отступа в точках


    def makeSubtitle(self, text):
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run(text)
        font = run.font
        font.size = Pt(16)
        font.name = 'IBM Plex Sans'  # Set the font name
        font.bold = True

    def makeTermsSection(self, terms):
        self.makeHeader("Термины, используемые в лекции")

        for item in terms:
            paragraph = self.doc.add_paragraph()
            run = paragraph.add_run(item[0] + " ")
            font = run.font
            font.size = Pt(12)
            font.name = 'IBM Plex Sans'  # Set the font name
            font.bold = True

            run = paragraph.add_run("- " + item[1])
            font = run.font
            font.size = Pt(12)
            font.name = 'IBM Plex Sans'  # Set the font name
            font.bold = False



    def makeHeader(self, text):
        paragraph = self.doc.add_heading()
        run = paragraph.add_run(text)
        font = run.font
        font.size = Pt(22)
        font.name = 'IBM Plex Sans'  # Set the font name
        font.bold = True

    def makeTextWithBoldWords(self, words, paragraph):
        for item in words:
            run = paragraph.add_run(item[0])
            font = run.font
            font.size = Pt(12)  # Set the font size
            font.name = 'IBM Plex Sans'  # Set the font name
            font.bold = item[1]

    def setBackgroundImage(self, image_path):
        # Add a section to the document


        self.doc.sections[0].left_margin = Inches(0.3)
        self.doc.add_picture(image_path, width=Inches(8.0))

        new_section = self.doc.add_section()
        new_section.left_margin = Inches(1.0)
        # run = self.doc.add_run()
        # run.add_picture(image_path, width=Inches(8))

    def makeLectureDocs(self, text, terms, image_path, title, subject):

        self.createTitlePage(title, subject, "doc_back.png")
        self.doc.add_page_break()
        self.makeHeader("Оглавление")
        shortContests = [[item[0], "0"] for item in text]
        self.makeShortContent(shortContests)
        self.makeTermsSection(terms=terms)
        for section in text:
            self.doc.add_page_break()
            self.makeHeader(section[0])
            paragraph = self.doc.add_paragraph()
            self.makeTextWithBoldWords(section[1], paragraph)

    def setPageBackground(self, image_path):
        section = self.doc.sections[0]
        new_width, new_height = section.page_width, section.page_height
        image = self.doc.add_picture(image_path, width=new_width, height=new_height)

        # Set the image behind the text as a watermark
        shape = image._element
        shape.getparent().remove(shape)
        header = section.header
        paragraph = header.paragraphs[0]
        run = paragraph.add_run()
        run._element.append(shape)

    def saveDocx(self, id):
        self.doc.save(f"uploaded_docx/{id}.docx")

    def convertDocxToPdf(sef, id):
        convert(f"uploaded_docx/{id}.docx", f"uploaded_docx/{id}.pdf")

